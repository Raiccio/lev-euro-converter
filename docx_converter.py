"""
DOCX converter for BGN to EUR
Uses python-docx for reliable DOCX handling
"""

import os
import re
from docx import Document
from converter_numbers import convert_bgn_to_eur, format_eur_words, parse_bulgarian_number, word_to_number

RATE = 1.95583


def convert_docx(input_path, output_path):
    """Convert DOCX file from BGN to EUR
    
    Returns:
        dict: {'success': bool, 'changes': list of str, 'error': str or None}
    """
    
    changes = []
    try:
        doc = Document(input_path)
        
        # Process paragraphs and their runs
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph, changes)
        
        # Process table cells and their runs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        process_paragraph(paragraph, changes)
        
        doc.save(output_path)
        
        return {'success': True, 'changes': changes, 'error': None}
        
    except Exception as e:
        return {'success': False, 'changes': changes, 'error': str(e)}


def process_paragraph(paragraph, changes):
    """Process a paragraph, preserving run formatting"""
    
    # Skip empty paragraphs
    if not paragraph.text or not paragraph.text.strip():
        return
    
    # Check if this paragraph needs conversion
    new_text, section_changes = process_text(paragraph.text)
    if not section_changes:
        return
    
    changes.extend(section_changes)
    
    # Build new runs while preserving formatting
    # Strategy: Create new runs with the new text, preserving formatting from original runs
    
    # First, collect all runs' formatting
    runs_data = []
    for run in paragraph.runs:
        runs_data.append({
            'text': run.text,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'font_bold': run.font.bold,
            'font_italic': run.font.italic,
        })
    
    # Clear all runs
    for _ in range(len(paragraph.runs)):
        p = paragraph._element
        p.remove(paragraph.runs[0]._element)
    
    # Add new run with new text and preserved formatting
    if runs_data:
        # Use formatting from first run as base
        first_run = runs_data[0]
        
        run = paragraph.add_run(new_text)
        
        if first_run['font_name']:
            run.font.name = first_run['font_name']
        if first_run['font_size']:
            run.font.size = first_run['font_size']
        if first_run['font_bold']:
            run.font.bold = first_run['font_bold']
        if first_run['font_italic']:
            run.font.italic = first_run['font_italic']


def has_transliteration(text):
    """Check if text has word transliteration in / or ()"""
    if not text:
        return False
    return bool(re.search(r'/[а-яА-Я\w]+/|/[а-яА-Я\w]+\)', text, re.IGNORECASE))


def process_text(text):
    """Process text and replace BGN amounts with EUR
    
    Returns:
        tuple: (new_text, list_of_changes)
    """
    if not text or not text.strip():
        return text, []
    
    # Check if text has Bulgarian currency indicator - if not, don't convert
    # Use case-insensitive substring matching
    currency_indicators = [
        'лева', 'лев', 'лв.', 'лв', 'lv.', 'lv', 'LV', 
        'BGN', 'leva', 'lev', 'български'
    ]
    text_lower = text.lower()
    has_bgn_currency = any(ind.lower() in text_lower for ind in currency_indicators)
    
    if not has_bgn_currency:
        return text, []
    
    # Clean up \xa0 (non-breaking space) first
    text = text.replace('\xa0', ' ')
    text = text.replace('\u00a0', ' ')
    
    original_text = text
    changes = []
    
    # Check if transliteration exists in this text (any /.../ or (...))
    text_has_transliteration = has_transliteration(text)
    
    # Define all currency variations (lowercase + uppercase handled by regex)
    currencies = r'лева|лев|лв\.?|лв|lv\.?|lv|LV|BGN|leva|lev|български\s*лева'
    
    # Pattern 1a: Currency AFTER slash: "50 /петдесет/ лева" or "100 000,00 /сто хиляди/ лв."
    slash_pattern_after = rf'(\d[\d\s]*(?:[.,]\d+)?)\s*/([а-яА-Я\w\s]+)/\s*(?:{currencies})?'
    for match in re.findall(slash_pattern_after, text, re.IGNORECASE):
        number_str = match[0].strip()
        word_part = match[1].strip()
        
        amount = parse_bulgarian_number(number_str)
        if amount is None:
            continue
        
        eur = convert_bgn_to_eur(amount)
        eur_words = format_eur_words(eur)
        output = f"{eur:.2f} /{eur_words}/ евро"
        
        # Try various currency suffixes
        found = False
        for suffix in [' лева', ' лев', ' лв.', ' лв', ' LV', ' LV.', ' lv.', ' lv', ' BGN', ' leva', ' LEVA', ' ЛЕВА', ' ЛЕВ', ' ЛВ.', ' ЛВ', '']:
            search = f"{number_str} /{word_part}/{suffix}"
            if search in text:
                changes.append(f"{search} --> {output}")
                text = text.replace(search, output)
                found = True
                break
        if not found:
            for suffix in ['лева', 'лев', 'лв.', 'лв', 'LV', 'LV.', 'lv.', 'lv', 'BGN', 'leva', 'LEVA', 'ЛЕВА', 'ЛЕВ', 'Л�.', 'ЛВ', 'LEV']:
                search = f"{number_str} /{word_part}/{suffix}"
                if search in text:
                    changes.append(f"{search} --> {output}")
                    text = text.replace(search, output)
                    found = True
                    break
    
    # Pattern 1b: Currency BEFORE slash: "50 лв /петдесет/" or "100 000,00 lev /сто хиляди/"
    slash_pattern_before = rf'(\d[\d\s]*(?:[.,]\d+)?)\s*({currencies})\s*/([а-яА-Я\w\s]+)/'
    for match in re.findall(slash_pattern_before, text, re.IGNORECASE):
        number_str = match[0].strip()
        currency = match[1].strip()
        word_part = match[2].strip()
        
        amount = parse_bulgarian_number(number_str)
        if amount is None:
            continue
        
        eur = convert_bgn_to_eur(amount)
        eur_words = format_eur_words(eur)
        output = f"{eur:.2f} /{eur_words}/ евро"
        
        search = f"{number_str} {currency} /{word_part}/"
        if search in text:
            changes.append(f"{search} --> {output}")
            text = text.replace(search, output)
    
    # Pattern 2a: Currency AFTER paren: "50 (петдесет) лева" or "100 000,00 (сто хиляди) лв."
    paren_pattern_after = rf'(\d[\d\s]*(?:[.,]\d+)?)\s*\(([а-яА-Я\w\s]+)\)\s*(?:{currencies})?'
    for match in re.findall(paren_pattern_after, text, re.IGNORECASE):
        number_str = match[0].strip()
        
        amount = parse_bulgarian_number(number_str)
        if amount is None:
            continue
        
        eur = convert_bgn_to_eur(amount)
        eur_words = format_eur_words(eur)
        output = f"{eur:.2f} ({eur_words}) евро"
        
        # Try various currency suffixes
        found = False
        for suffix in [' лева', ' лев', ' лв.', ' лв', ' LV', ' LV.', ' lv.', ' lv', ' BGN', ' leva', '']:
            search = f"{number_str} ({match[1]}){suffix}"
            if search in text:
                changes.append(f"{search} --> {output}")
                text = text.replace(search, output)
                found = True
                break
        if not found:
            for suffix in ['лева', 'лев', 'лв.', 'лв', 'LV', 'LV.', 'lv.', 'lv', 'BGN', 'leva']:
                search = f"{number_str} ({match[1]}){suffix}"
                if search in text:
                    changes.append(f"{search} --> {output}")
                    text = text.replace(search, output)
                    found = True
                    break
    
    # Pattern 2b: Currency BEFORE paren: "50 лв (петдесет)" or "100 000,00 lev (сто хиляди)"
    paren_pattern_before = rf'(\d[\d\s]*(?:[.,]\d+)?)\s*({currencies})\s*\(([а-яА-Я\w\s]+)\)'
    for match in re.findall(paren_pattern_before, text, re.IGNORECASE):
        number_str = match[0].strip()
        
        amount = parse_bulgarian_number(number_str)
        if amount is None:
            continue
        
        eur = convert_bgn_to_eur(amount)
        eur_words = format_eur_words(eur)
        output = f"{eur:.2f} ({eur_words}) евро"
        
        search = f"{number_str} {match[1]} ({match[2]})"
        if search in text:
            changes.append(f"{search} --> {output}")
            text = text.replace(search, output)
    
    # Pattern 3: Simple currency (no words) - "100 лв." or "100 лева" or "100,00 BGN"
    currency_patterns = [
        r'(\d[\d\s]*(?:[.,]\d+)?)\s*(лв\.|лв|лева|лев|lv\.?|lv|LV|BGN|български\s*лева|leva|LEV|LEVA|ЛЕВА|ЛЕВ|ЛВ\.|ЛВ)',
        r'(\d[\d\s]*(?:[.,]\d+)?)\s*(стотинки|ст\.)',
    ]
    
    for pattern in currency_patterns:
        for match in re.findall(pattern, text, re.IGNORECASE):
            amount_str = match[0].strip()
            currency = match[1]
            
            amount = parse_bulgarian_number(amount_str)
            if amount is None:
                continue
            
            eur = convert_bgn_to_eur(amount)
            
            # Only use Bulgarian words if text has transliteration elsewhere
            if text_has_transliteration:
                output = format_eur_words(eur)
            else:
                # Just output the number
                if currency.lower() in ['стотинки', 'ст.']:
                    output = f"{eur:.2f} евроцента"
                else:
                    output = f"{eur:.2f} евро"
            
            # Try different spacing
            search = f"{amount_str} {currency}"
            if search in text:
                changes.append(f"{search} --> {output}")
                text = text.replace(search, output)
            elif f"{amount_str}  {currency}" in text:
                text = text.replace(f"{amount_str}  {currency}", output)
            elif f"{amount_str}{currency}" in text:
                text = text.replace(f"{amount_str}{currency}", output)
    
    return text, changes


def get_output_path(input_path):
    """Generate output path"""
    dir_name = os.path.dirname(input_path)
    base_name = os.path.basename(input_path)
    name, ext = os.path.splitext(base_name)
    return os.path.join(dir_name, f"{name}_eur{ext}")